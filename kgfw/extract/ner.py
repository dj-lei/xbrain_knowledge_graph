import spacy
import pandas as pd
from abc import abstractmethod
from docx import Document


class BaseNer(object):

    def __init__(self, model_name, path):
        self.proper_nouns = []
        self.nlp = spacy.load(model_name)
        self.document = Document(path)

    def get_stop_words(self, doc):
        result = []
        for token in doc:
            if token.is_stop:
                result.append(token.text)
        return result

    def get_verb(self, doc):
        result = []
        for token in doc:
            if token.pos_ in ['VERB']:
                result.append(token.lemma_)
        return result

    def get_noun_chunks(self, doc):
        stop_words = self.get_stop_words(doc)
        result = []
        for chunk in doc.noun_chunks:
            chunk = chunk.text
            result.append(' '.join(['' if t in stop_words else t for t in chunk.split(' ')]).strip().replace('  ',' '))
        return result

    def get_text_token_pos(self, text, token):
        for i in range(0, len(text)):
            if text[i:i + len(token)] == token:
                return (i, i + len(token) - 1)
        return None

    @abstractmethod
    def get_proper_nouns(self):
        pass

    def get_rules(self):
        pass

    def model(self):
        pass

    def train(self):
        pass

    def eval(self):
        pass


class RuNer(BaseNer):

    def __init__(self, model_name, path):
        super(RuNer, self).__init__(model_name, path)
        self.get_proper_nouns()

    def get_proper_nouns(self):
        """
        get proper nouns
        """
        self.proper_nouns = list(pd.read_csv('proper_nouns.csv')['proper_nouns_name'].values)

    def get_sentences_ner(self):
        count = 0
        data = pd.DataFrame([], columns=['style_name', 'text', 'ner'])
        for i in range(0, len(self.document.paragraphs)):
            doc = self.document.paragraphs[i].text.strip()
            if doc not in ['', '\n']:
                for sentence in doc.split('.'):
                    if (len(sentence.strip().split(' ')) <= 3) | (self.document.paragraphs[i].style.name in ['Caption']):
                        continue
                    if len(set(self.proper_nouns).intersection(set(sentence.split(' ')))) == 0:
                        continue

                    noun_chunks = self.get_noun_chunks(self.nlp(sentence))
                    data.loc[count, 'style_name'] = self.document.paragraphs[i].style.name
                    data.loc[count, 'text'] = sentence
                    data.loc[count, 'ner'] = str(noun_chunks)
                    count = count + 1
        data['file_name'] = self.document.core_properties.title
        data = data.reset_index(drop=True)

        return data

